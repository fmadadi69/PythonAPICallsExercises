import docx
filename = '/Users/apple/Documents/intro.docx'
new_filename= '/Users/apple/Documents/new_intro.docx'

###################### Method 1 #######################
# def getText(text_file):
#     doc = docx.Document(text_file)
#     fulltext = []
#     for para in doc.paragraphs:
#         fulltext.append(para.text)
#     return '\n'.join(fulltext)
#
#
# text = getText(filename)
#
# all_nums = re.findall('\d+', text)
# all_reversed_nums = []
# for nums in all_nums:
#     all_reversed_nums.append(nums[::-1])
#
# print(all_nums)
# print(all_reversed_nums)
#
# new_text = text
#
# # for item1 in all_nums:
# #     new_text = new_text.replace(item1, all_reversed_nums.pop(0))
#
# zipped_nums = zip(all_nums, all_reversed_nums)
# for items in zipped_nums:
#     new_text = new_text.replace(items[0],items[1])
#
# print(new_text)
# my_doc = docx.Document()
# my_doc.add_paragraph(new_text)
# my_doc.save(new_filename)

#################### Method2 ######################
doc = docx.Document(filename)
numbers_done = 0

for para in doc.paragraphs:
    new_text = ''
    for char in para.text:
        if char in '0123456789':
            if numbers_done == 0:
                new_text += char
                numbers_done += 1
            else:
                new_text = new_text[:-1 * numbers_done] + char + new_text[-1 * numbers_done:]
                numbers_done += 1
        else:
            new_text += char
            numbers_done = 0
    para.text = new_text
doc.save(new_filename)


